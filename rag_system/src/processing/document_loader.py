"""
Document Loader
Supports PDF, TXT, Markdown, DOCX
"""

import os
from pathlib import Path
from typing import Dict, List, Optional
from dataclasses import dataclass
import hashlib

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

import markdown
from bs4 import BeautifulSoup


@dataclass
class Document:
    """Document data class"""
    content: str
    metadata: Dict
    doc_id: str
    source: str


class DocumentLoader:
    """
    Load documents from various formats

    Supported formats:
    - PDF (.pdf)
    - Text (.txt)
    - Markdown (.md)
    - Word (.docx)
    """

    def __init__(self):
        self.supported_extensions = {'.pdf', '.txt', '.md', '.markdown', '.docx'}

    def load_file(self, file_path: str) -> Optional[Document]:
        """
        Load a single file

        Args:
            file_path: Path to file

        Returns:
            Document object or None if unsupported
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = path.suffix.lower()

        if extension not in self.supported_extensions:
            print(f"⚠️  Unsupported file type: {extension}")
            return None

        # Extract content based on file type
        if extension == '.pdf':
            content = self._load_pdf(path)
        elif extension in {'.txt'}:
            content = self._load_text(path)
        elif extension in {'.md', '.markdown'}:
            content = self._load_markdown(path)
        elif extension == '.docx':
            content = self._load_docx(path)
        else:
            return None

        if not content or not content.strip():
            print(f"⚠️  No content extracted from: {file_path}")
            return None

        # Generate document ID
        doc_id = self._generate_doc_id(content, str(path))

        # Extract metadata
        metadata = self._extract_metadata(path, content)

        return Document(
            content=content,
            metadata=metadata,
            doc_id=doc_id,
            source=str(path)
        )

    def load_directory(self, directory_path: str, recursive: bool = True) -> List[Document]:
        """
        Load all supported documents from directory

        Args:
            directory_path: Path to directory
            recursive: Recursively search subdirectories

        Returns:
            List of Document objects
        """
        path = Path(directory_path)

        if not path.exists() or not path.is_dir():
            raise ValueError(f"Invalid directory: {directory_path}")

        documents = []

        # Find all files
        if recursive:
            files = path.rglob('*')
        else:
            files = path.glob('*')

        for file_path in files:
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    doc = self.load_file(str(file_path))
                    if doc:
                        documents.append(doc)
                except Exception as e:
                    print(f"❌ Error loading {file_path}: {e}")

        print(f"✅ Loaded {len(documents)} documents from {directory_path}")
        return documents

    def _load_pdf(self, path: Path) -> str:
        """Load PDF file"""
        if PdfReader is None:
            raise ImportError("pypdf is required for PDF support. Install with: pip install pypdf")

        try:
            reader = PdfReader(str(path))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            print(f"❌ Error reading PDF {path}: {e}")
            return ""

    def _load_text(self, path: Path) -> str:
        """Load text file"""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(path, 'r', encoding=encoding) as f:
                        return f.read()
                except:
                    continue
            print(f"❌ Could not decode {path}")
            return ""

    def _load_markdown(self, path: Path) -> str:
        """Load markdown file and convert to plain text"""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                md_content = f.read()

            # Convert markdown to HTML
            html = markdown.markdown(md_content)

            # Extract text from HTML
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text()

            return text.strip()
        except Exception as e:
            print(f"❌ Error reading Markdown {path}: {e}")
            return ""

    def _load_docx(self, path: Path) -> str:
        """Load Word document"""
        if DocxDocument is None:
            raise ImportError("python-docx is required for DOCX support. Install with: pip install python-docx")

        try:
            doc = DocxDocument(str(path))
            text = "\n".join([para.text for para in doc.paragraphs])
            return text.strip()
        except Exception as e:
            print(f"❌ Error reading DOCX {path}: {e}")
            return ""

    def _generate_doc_id(self, content: str, source: str) -> str:
        """Generate unique document ID based on content hash"""
        content_hash = hashlib.sha256(content.encode('utf-8')).hexdigest()[:16]
        source_hash = hashlib.sha256(source.encode('utf-8')).hexdigest()[:8]
        return f"{source_hash}_{content_hash}"

    def _extract_metadata(self, path: Path, content: str) -> Dict:
        """Extract metadata from document"""
        return {
            'filename': path.name,
            'extension': path.suffix,
            'size_bytes': path.stat().st_size,
            'size_chars': len(content),
            'size_words': len(content.split()),
            'created': path.stat().st_ctime,
            'modified': path.stat().st_mtime,
        }
